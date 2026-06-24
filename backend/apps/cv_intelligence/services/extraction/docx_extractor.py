"""Extract structured text from DOCX using python-docx."""

from __future__ import annotations

import io
import re


def extract_text_from_docx(file_bytes: bytes) -> tuple[str, dict]:
    metadata: dict = {'extractors': []}
    if not file_bytes:
        return '', metadata

    try:
        from docx import Document
    except ImportError:
        metadata['docx_error'] = 'not_installed'
        return file_bytes.decode('utf-8', errors='ignore'), metadata

    try:
        doc = Document(io.BytesIO(file_bytes))
        parts: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(' | '.join(cells))
        metadata['extractors'].append('python-docx')
        metadata['paragraph_count'] = len(doc.paragraphs)
        metadata['table_count'] = len(doc.tables)
        return '\n'.join(parts), metadata
    except Exception as exc:
        metadata['docx_error'] = str(exc)
        return file_bytes.decode('utf-8', errors='ignore'), metadata


def extract_hyperlinks_from_docx(file_bytes: bytes) -> dict[str, str]:
    """Best-effort link extraction from DOCX relationships."""
    links: dict[str, str] = {}
    try:
        from docx import Document
    except ImportError:
        return links

    try:
        doc = Document(io.BytesIO(file_bytes))
        for rel in doc.part.rels.values():
            if 'hyperlink' in rel.reltype:
                target = rel.target_ref
                if target:
                    links[target] = target
    except Exception:
        pass
    return links
